import pandas as pd
import numpy as np
import os
import docx
from docx.shared import Cm
from circuit import Circuit
from operator import itemgetter
from itertools import groupby


def start_update(circuit):
    IDs = circuit.set_IDs

    # Start Updates
    make_boxmaps(circuit, IDs)

def make_boxmaps(circuit, IDs):
    # Drug_Type_Box#_Date
    files = get_filenames("../../Box Maps")
    circuit_drugs = circuit.drugs.values()
    circuit_type = circuit.type
    existing_maps = []
    filtered = []
    needed = []

    # Go through each box
    for type, cirs in groupby(sorted(IDs), key=itemgetter(0)):
        for drug, circuits in groupby(sorted(cirs), lambda x : x[2:4]):
            df = pd.DataFrame(columns = ["Box Label",
                                            "Cell",
                                            "Sample ID",
                                            "Date",
                                            "Timepoint"])
            cur_type = 'ECMO' if type == 'E' else 'CRRT'
            cur_drug = circuit.drugs[drug]
            cand_box = None
            box_num = 1
            num_samples = circuit.get_samples_per_drug()
            file_date = circuit.file_date
            tube_date = circuit.tube_date
            cir_names = list(circuits)
            cells = generate_cells()
            start = 0
            new = True

            # look up room in current boxes
            # Get last used box
            for f in files:
                name = f.split("_")
                if name[0] == cur_drug and name[1] == cur_type:
                    if box_num <= int(name[2][3:]):
                        box_num = int(name[2][3:])
                        cand_box = f

            # Check candidate box for room or make new box
            if cand_box:
                box = pd.read_csv("../../Box Maps/" + cand_box)
                box = box[box['Sample ID'] != 'Empty']
                room = 81 - len(box)
                if room >= num_samples:
                    start = 81 - room
                    label = cand_box[:-4]
                    new = False
                else:
                    box_num = box_num + 1
            # Check if new box needs to be made
            if new:
                cand_box = "{}_{}_Box{}_{}.csv".format(cur_drug,
                                                    cur_type,
                                                    box_num,
                                                    file_date)
                label = "{} - {} Box {}, Watt {}".format(cur_drug,
                                                    cur_type,
                                                    box_num,
                                                    file_date[-4:])
            # Build dataframe
            time = 0
            count = range(start, 81)
            fill = True
            mod = len(cir_names)
            for idx, x in enumerate(count):
                if idx%mod == 0 and fill:
                    time = time + 1
                _type = cir_names[idx%mod]
                if idx >= num_samples:
                    sample_id = "Empty"
                    tube_date = ""
                    timepoint = ""
                    fill = False
                else:
                     sample_id = _type + str(time)
                     timepoint = circuit.timepoints[time]
                df.loc[len(df)] = [label,
                                    cells[x],
                                    sample_id,
                                    tube_date,
                                    timepoint]
            if not new:
                df = pd.concat([box, df])

            df.to_csv("../../Box Maps/" + cand_box, index=False)

def build_labs_crf(circuit, drug):
    doc = docx.Document("../../CRF Templates and Labels/Labs_CRF.docx")
    drug_long = circuit.drugs[drug]
    table = doc.tables[0]
    for i in range(1,11):
        table.cell(i,0).text = circuit.type
        table.cell(i,1).text = circuit.circuit_num
        table.cell(i,2).text = drug
        table.cell(i,3).text = circuit.tube_date

        if i + 3 <= circuit.samples_per_type and i != 1:
            table.cell(i,5).text = str(i + 3)
        elif i == 1:
            table.cell(i,5).text = "0"

    file = build_file_name(circuit, "docx", drug_long, other="_Labs_CRF")
    doc.save("../../CRF Templates and Labels/" + file)

def build_sampling_crf(circuit, drug):
    doc = docx.Document("../../CRF Templates and Labels/CRF_Sample_Collection_Template.docx")
    table = doc.tables[0]
    num_rows = circuit.get_samples_per_drug()
    extra_rows = 1
    sample_types = circuit.get_sample_types()
    mod = len(sample_types)
    drug_long = circuit.drugs[drug]

    if circuit.control:
        row = table.add_row()
        extra_rows = 2
        for i in range(2,6):
            cell = table.cell(2,1)
            cell.merge(table.cell(2,i))
        table.cell(1,1).text = "DOSE 1 CONTROL"
        table.cell(1,6).text = "-5m"
        table.cell(2,1).text = "DOSE 1 CIRCUIT"
        table.cell(2,6).text = "0m"

    s_num = 0
    for i in range(num_rows):
        row = table.add_row()
        site = sample_types[i%mod]
        cells = row.cells
        if i%mod == 0:
            s_num = s_num+1
        cells[0].text = circuit.type
        cells[1].text = circuit.circuit_num
        cells[2].text = drug
        cells[3].text = site
        cells[4].text = str(s_num)
        cells[5].text = circuit.tube_date
        cells[6].text = circuit.timepoints[s_num]

    for row in table.rows:
        row.height = Cm(1.1)
    file = build_file_name(circuit, "docx", drug_long, other="_Samples_CRF")
    doc.save("../../CRF Templates and Labels/" + file)


    doc.save("test.docx")

def build_file_name(circuit, extension, drug, other=""):
    date = circuit.file_date
    type = circuit.type

    return "{}_{}_{}{}.{}".format(drug, type, date, other, extension)

def get_filenames(path):
    files = [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
    return files

def generate_cells():
    rows = ['A','B','C','D','E','F','G','H','I']
    cols = range(1,10)
    cells = []
    for r in rows:
        for c in cols:
            cells.append(r+str(c))
    return cells

def main():
    test = Circuit("2",
                    "6 hr",
                    {1:"1 min", 2:"5 min", 3:"15 min", 4:"30 min", 5:"1 hr",
                     6:"2 hr", 7:"3 hr", 8:"4 hr", 9:"5 hr", 10:"6 hr", 11:"8 hr"},
                     {"Dx":"Dexmedtomidine", "Cl":"Clindamycin"},
                     "CRRT",
                     "True",
                     "2022-6-23")

    build_sampling_crf(test, "Dx")

if __name__ == "__main__":
    main()
